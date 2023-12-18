import asyncio
from selectolax.parser import HTMLParser
import aiohttp
import qrcode
from docx import Document
from docx.shared import Pt

async def fetch(url:str) -> str:
    async with aiohttp.ClientSession() as session:
        async with session.get(url) as response:
            return await response.text()
        
async def parse_catalog() -> list:
    html = await fetch('https://www.boyard.biz/catalog')
    parser = HTMLParser(html)
    links = [a.attrs['href'] for a in parser.css('.product-card')]
    return links

async def parse_links(link:str) -> str:
    html = await fetch(link)
    parser = HTMLParser(html)
    links = {}
    for a in parser.css('.bd-product__link'):
        links.update({a.attributes['title']: a.attributes['href']})
    return links

async def make_qr(text: str, name:str) -> None:
    qr = qrcode.make(text)
    name = text.split('/')[-1:][0]
    qr.save('qr/'+name+'.png')

async def save_data(data):
    from model import DataSaver
    data_saver = DataSaver('data.db')
    data_saver.create_table('qr', ['name TEXT', 'url TEXT'])
    for item in data:
        data_saver.insert_data('qr', item)
    data_saver.close_connection()

async def main():
    catalog = await parse_catalog()
    doc = Document()
    doc.add_heading('QR-коды продуктов Boyard', 0)
    for link in catalog:
        links = await parse_links(link)
        for k in links:
            name = links[k].split('/')[-1:][0][:-5]
            await make_qr(links[k], name)
            doc.add_paragraph(k)
            doc.add_picture('qr/'+name+'.png', width=Pt(100))
            print(name+'.png')
    doc.save('catalog.docx')

if __name__ == "__main__":
    loop = asyncio.get_event_loop()
    loop.run_until_complete(main())