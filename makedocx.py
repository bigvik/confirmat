import math
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from model import DataSaver




def get_data():
    '''
    Return id, url from product_data
    '''
    data_saver = DataSaver('confirmat.db')
    data = data_saver.get_urls_product()
    data_saver.close_connection()
    return data

def get_datawhere():
    filter = ['DB8881', 'DB8882', 'SB08', 'H301', 'H306', 'H642', 'RT109', 'RT110', 'RT111', 'RP050', 'RP051']
    data_saver = DataSaver('confirmat.db')
    data = []
    for f in filter:
        dw = data_saver.get_where(f)
        data = data + dw
    data_saver.close_connection()
    return data

def make_doc():
    pass

def make_docwhere():
    data = get_datawhere()
    doc = Document()
    doc.add_heading('QR Каталог: с фильтром', 0)
    counter = 1
    l = len(data)
    for d in data:
        print(f'Processing {counter}/{l}', end='\r')
        name = d[0].split('/')[-1:][0].replace('*', '_')
        path = 'qr/'+name+'.png'
        doc.add_paragraph(d[1])
        doc.add_picture(path, width=Pt(100))
        counter += 1
    doc.save('QRwithFilter.docx')
    print('Done!')

def make_test():
    data = get_datawhere()
    doc = Document()
    doc.add_heading('QR Каталог: с фильтром', 0)

    data_len = len(data)
    rows = math.ceil(data_len / 4)
    counter = 1
    table = doc.add_table(rows=0, cols=4, style="Table Grid")

    for _ in range(rows):
        img_row = table.add_row()
        cap_row = table.add_row()
        for col in range(4):
            print(counter, end='\r')
            if len(data) == 0: break
            counter += 1
            img = data.pop()
            name = img[0].split('/')[-1:][0].replace('*', '_')
            path = 'qr/'+name+'.png'
            # add image to table
            #set_cell_margins(img_row.cells[col], top=100, start=100, bottom=100, end=50)
            # add image to cell and align center
            paragraph = img_row.cells[col].paragraphs[0] 
            paragraph.add_run().add_picture(path, width=Pt(70))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # add caption to table
            cap_row.cells[col].text = name.split('.')[0]
            cap_row.cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # save doc
    doc.save('QRwithFilter3.docx')
    print('Done!')

def main():
    doc = Document()
    doc.add_heading('QR Каталог', 0)
    data = get_data()
    counter = 1
    l = len(data)
    for d in data:
        print(f'Processing {counter}/{l}', end='\r')
        name = d[0].split('/')[-1:][0].replace('*', '_')
        path = 'qr/'+name+'.png'
        doc.add_paragraph(d[1])
        doc.add_picture(path, width=Pt(100))
        counter += 1
    doc.save('QRcatalog.docx')
    print('Done!')


if __name__ == '__main__':
    #main()
    #make_docwhere()
    make_test()